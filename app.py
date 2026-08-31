import os
import subprocess
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd
import requests
import streamlit as st
from docx import Document

from auth import cookie_manager, logout, require_login
from gasu import (
    GASU_FOUNDED_YEAR,
    apply_author_profile,
    author_id_query,
    author_ids,
    author_papers_query,
    author_profile_query,
    author_search_id,
    build_query,
    entry_belongs_to_gasu,
    first_initial,
    format_affiliations,
    needs_author_enrichment,
    parse_author_retrieval,
    parse_author_search_profile,
    parse_authors,
    pick_scopus_authid,
    query_targets_gasu,
)
from report import (
    ReportData,
    build_report,
    report_area_png,
    report_chart_png,
    report_quartile_png,
    report_scope_label,
    report_sentence,
    rsf_candidates,
    rsf_eligibility_rows,
    ru_publications,
    top_authors,
    apply_author_total,
)
from rsf import record_in_rsf_window, rsf_window
from scimago import (
    attach_scimago,
    format_quartile_cell,
    format_ru_date,
    load_meta,
    lookup_built_on,
    next_refresh_date,
    quartile_share_rows,
)
from subjects import (
    area_share_rows,
    attach_subject_areas,
    extract_issns,
    fetch_serial_abbrevs,
    grouped_area_counts,
)

try:
    from dotenv import load_dotenv
except Exception:
    load_dotenv = None

API_URL = "https://api.elsevier.com/content/search/scopus"
AUTHOR_URL = "https://api.elsevier.com/content/search/author"
AUTHOR_RETRIEVAL_URL = "https://api.elsevier.com/content/author/author_id"
ABSTRACT_URL = "https://api.elsevier.com/content/abstract"
SEARCH_FIELDS = (
    "dc:identifier,dc:title,dc:creator,prism:coverDate,prism:publicationName,"
    "prism:doi,prism:volume,prism:issueIdentifier,prism:pageRange,"
    "prism:issn,prism:eIssn,author,affiliation"
)
ENV_PATH = Path(__file__).with_name(".env")
APP_VERSION = "1.8.7"
APP_UPDATED_FALLBACK = "29.08.2026"


def last_updated_label() -> str:
    try:
        raw = subprocess.check_output(
            ["git", "log", "-1", "--format=%cs"],
            cwd=Path(__file__).resolve().parent,
            stderr=subprocess.DEVNULL,
            timeout=2,
            text=True,
        ).strip()
        if raw:
            return datetime.strptime(raw, "%Y-%m-%d").strftime("%d.%m.%Y")
    except Exception:
        pass
    return APP_UPDATED_FALLBACK


def running_on_streamlit_cloud() -> bool:
    return Path("/mount/src").exists() or bool(os.getenv("STREAMLIT_SHARING_MODE"))


def load_api_key() -> str | None:
    try:
        if "SCOPUS_API_KEY" in st.secrets:
            return st.secrets["SCOPUS_API_KEY"]
    except Exception:
        pass
    if load_dotenv:
        load_dotenv(ENV_PATH)
    return os.getenv("SCOPUS_API_KEY")


def save_api_key(value: str) -> None:
    value = value.strip()
    if not value:
        return
    lines = []
    if ENV_PATH.exists():
        lines = ENV_PATH.read_text(encoding="utf-8").splitlines()
    key_line = f"SCOPUS_API_KEY={value}"
    updated = False
    for i, line in enumerate(lines):
        if line.startswith("SCOPUS_API_KEY="):
            lines[i] = key_line
            updated = True
            break
    if not updated:
        lines.append(key_line)
    ENV_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")
    os.environ["SCOPUS_API_KEY"] = value


def normalize_initials(text: str) -> str:
    cleaned = (text or "").replace(".", "").replace("-", " ").strip()
    if not cleaned:
        return ""
    parts = [p for p in cleaned.split() if p]
    return "".join(f"{p[0].upper()}." for p in parts)


def initials_from_given(given: str) -> str:
    return normalize_initials(given)


def format_authors_gost(authors: list[dict]) -> str:
    formatted = []
    for author in authors:
        surname = author.get("surname", "").strip()
        given = author.get("given", "").strip()
        initials = initials_from_given(given) or normalize_initials(author.get("initials", ""))
        if surname and initials:
            formatted.append(f"{surname} {initials}")
        elif surname:
            formatted.append(surname)
    return ", ".join(formatted)


def format_authors_apa(authors: list[dict]) -> str:
    formatted = []
    for author in authors:
        surname = author.get("surname", "").strip()
        given = author.get("given", "").strip()
        initials = initials_from_given(given) or normalize_initials(author.get("initials", ""))
        if surname and initials:
            formatted.append(f"{surname}, {initials}")
        elif surname:
            formatted.append(surname)
    if not formatted:
        return ""
    if len(formatted) == 1:
        return formatted[0]
    return ", ".join(formatted[:-1]) + f", & {formatted[-1]}"


def format_gost(record: dict) -> str:
    parts = []
    authors = format_authors_gost(record["authors"])
    if authors:
        parts.append(authors)
    if record["title"]:
        parts.append(record["title"])
    main = " ".join(parts).strip()
    journal_part = f"// {record['journal']}" if record["journal"] else ""
    year_part = f"{record['year']}" if record["year"] else ""
    volume_part = f"Т. {record['volume']}" if record["volume"] else ""
    issue_part = f"№ {record['issue']}" if record["issue"] else ""
    pages_part = f"С. {record['pages']}" if record["pages"] else ""
    tail = ". ".join([p for p in [journal_part, year_part, volume_part, issue_part, pages_part] if p])
    if tail:
        return f"{main} {tail}."
    return f"{main}."


def format_apa(record: dict) -> str:
    authors = format_authors_apa(record["authors"])
    year_part = f"({record['year']})." if record["year"] else "(n.d.)."
    title_part = f"{record['title']}." if record["title"] else ""
    journal_part = record["journal"] or ""
    volume_issue = ""
    if record["volume"] and record["issue"]:
        volume_issue = f"{record['volume']}({record['issue']})"
    elif record["volume"]:
        volume_issue = record["volume"]
    pages_part = record["pages"]
    doi = record["doi"]
    doi_part = f"https://doi.org/{doi}" if doi else ""
    tail = ", ".join([p for p in [journal_part, volume_issue, pages_part] if p])
    if tail:
        tail = f"{tail}."
    parts = [p for p in [authors, year_part, title_part, tail, doi_part] if p]
    return " ".join(parts).strip()


def make_date_filter(mode: str, start_year: int | None, end_year: int | None) -> dict | None:
    current_year = datetime.now().year
    if mode == "current":
        return {"mode": "current", "year": current_year, "year_start": current_year, "year_end": current_year}
    if mode == "last5":
        start = current_year - 4
        return {"mode": "range", "year": current_year, "year_start": start, "year_end": current_year}
    if mode == "since_founding":
        return {
            "mode": "range",
            "year": current_year,
            "year_start": GASU_FOUNDED_YEAR,
            "year_end": current_year,
        }
    if mode == "range" and start_year and end_year:
        return {"mode": "range", "year": start_year, "year_start": start_year, "year_end": end_year}
    return None


@st.cache_data(ttl=60 * 60 * 24 * 30, show_spinner=False)
def cached_serial_abbrevs(issn: str, api_key: str) -> tuple[list[str], int]:
    return fetch_serial_abbrevs(issn, api_key)


@st.cache_resource
def journal_memory() -> dict:
    return {"issn": {}}


def fill_subject_areas(records: list[dict], api_key: str) -> None:
    memory = journal_memory()
    cache = st.session_state.setdefault("issn_subject_cache", {})
    cache.update(memory["issn"])
    attach_subject_areas(
        records,
        api_key,
        cache,
        fetch=cached_serial_abbrevs,
    )
    memory["issn"].update(cache)


def _scopus_get(headers: dict, params: dict) -> requests.Response:
    last_error = None
    response = None
    for attempt in range(3):
        try:
            response = requests.get(API_URL, headers=headers, params=params, timeout=60)
            last_error = None
            break
        except requests.RequestException as exc:
            last_error = exc
            time.sleep(1.5 * (attempt + 1))
    if last_error or response is None:
        raise RuntimeError(f"Scopus API timeout: {last_error}") from last_error
    return response


def fetch_scopus_total(query: str, api_key: str) -> int:
    """Сколько документов находит запрос — без выгрузки и без фильтра ГАГУ."""
    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    response = _scopus_get(headers, {"query": query, "count": 1, "start": 0})
    if response.status_code != 200:
        raise RuntimeError(response.text)
    payload = response.json()
    return int((payload.get("search-results") or {}).get("opensearch:totalResults") or 0)


def _author_search_entries(query: str, api_key: str) -> list[dict]:
    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    try:
        response = requests.get(
            AUTHOR_URL,
            headers=headers,
            params={"query": query, "count": 15, "start": 0},
            timeout=45,
        )
    except requests.RequestException:
        return []
    if response.status_code != 200:
        return []
    try:
        payload = response.json()
    except Exception:
        return []
    entries = (payload.get("search-results") or {}).get("entry") or []
    if isinstance(entries, dict):
        entries = [entries]
    return [entry for entry in entries if isinstance(entry, dict)]


def resolve_author_profile(surname: str, initials: str, given: str, api_key: str) -> dict:
    if not (surname or "").strip():
        return {}
    queries = [author_profile_query(surname, initials, given)]
    if first_initial(initials, given):
        queries.append(author_profile_query(surname, initials, given, with_initial=False))
    seen: set[str] = set()
    for query in queries:
        if query in seen:
            continue
        seen.add(query)
        entries = _author_search_entries(query, api_key)
        authid = pick_scopus_authid(
            entries,
            surname=surname,
            initials=initials,
            given=given,
        )
        if not authid:
            continue
        for entry in entries:
            if author_search_id(entry) == authid:
                profile = parse_author_search_profile(entry)
                if profile.get("authid"):
                    return profile
        return {"authid": authid}
    return {}


def resolve_scopus_authid(surname: str, initials: str, given: str, api_key: str) -> str:
    return (resolve_author_profile(surname, initials, given, api_key).get("authid") or "").strip()


def _abstract_targets(record: dict) -> list[tuple[str, str]]:
    targets: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()

    def add(kind: str, ident: str) -> None:
        ident = (ident or "").strip()
        if not ident or (kind, ident) in seen:
            return
        seen.add((kind, ident))
        targets.append((kind, ident))

    sid = (record.get("scopus_id") or "").strip()
    eid = (record.get("eid") or "").strip()
    if sid.startswith("2-s2.0-"):
        add("eid", sid)
    elif sid:
        add("scopus_id", sid)
    add("eid", eid)
    return targets


def fetch_paper_authors(record: dict, api_key: str) -> tuple[list[dict], int]:
    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    last_status = 0
    for kind, ident in _abstract_targets(record):
        url = f"{ABSTRACT_URL}/{kind}/{ident}"
        for attempt in range(3):
            try:
                response = requests.get(url, headers=headers, timeout=45)
            except requests.RequestException:
                time.sleep(1.5 * (attempt + 1))
                continue
            last_status = response.status_code
            if response.status_code == 429:
                time.sleep(2 * (attempt + 1))
                continue
            if response.status_code != 200:
                break
            try:
                authors = parse_authors(response.json())
            except Exception:
                return [], last_status
            if authors:
                return authors, last_status
            break
    return [], last_status


def enrich_record_authors(records: list[dict], api_key: str) -> int:
    """По Scopus ID статьи добираем полный список авторов и их Author ID, если API открыт."""
    todo = [rec for rec in records if needs_author_enrichment(rec)]
    if not todo:
        return 0
    sample, status = fetch_paper_authors(todo[0], api_key)
    if status in {401, 403, 404} and not sample:
        return 0
    if sample:
        todo[0]["authors"] = sample
        todo = todo[1:]
    if not todo:
        return 1 if sample else 0
    updated = 1 if sample else 0
    with ThreadPoolExecutor(max_workers=4) as pool:
        futures = {pool.submit(fetch_paper_authors, rec, api_key): rec for rec in todo}
        for fut in as_completed(futures):
            rec = futures[fut]
            try:
                authors, _status = fut.result()
            except Exception:
                continue
            if authors:
                rec["authors"] = authors
                updated += 1
    return updated


def fetch_author_metrics(authid: str, api_key: str) -> dict:
    ident = "".join(ch for ch in str(authid or "") if ch.isdigit())
    if not ident:
        return {}
    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    params = {"view": "METRICS"}
    try:
        response = requests.get(
            f"{AUTHOR_RETRIEVAL_URL}/{ident}",
            headers=headers,
            params=params,
            timeout=45,
        )
    except requests.RequestException:
        return {}
    if response.status_code in {400, 401, 403}:
        try:
            response = requests.get(
                f"{AUTHOR_RETRIEVAL_URL}/{ident}",
                headers=headers,
                timeout=45,
            )
        except requests.RequestException:
            return {}
    if response.status_code != 200:
        return {}
    try:
        return parse_author_retrieval(response.json())
    except Exception:
        return {}


def stamp_author_profiles(records: list[dict], api_key: str) -> int:
    """Author Search + Retrieval: Author ID, ORCID, h-index и счётчики профиля."""
    groups: dict[str, list[dict]] = {}
    for rec in records:
        for author in rec.get("authors") or []:
            surname = (author.get("surname") or "").strip()
            if not surname:
                continue
            key = f"{surname.lower()}|{first_initial(author.get('initials') or '', author.get('given') or '').lower()}"
            groups.setdefault(key, []).append(author)
    if not groups:
        return 0

    def lookup(key: str) -> dict:
        sample = groups[key][0]
        if (sample.get("authid") or "").isdigit():
            return {
                "authid": sample.get("authid"),
                "orcid": sample.get("orcid") or "",
                "documents": sample.get("documents"),
                "cited_by": sample.get("cited_by"),
                "h_index": sample.get("h_index"),
                "profile_affil": sample.get("profile_affil") or "",
            }
        surname = (sample.get("surname") or "").strip()
        return resolve_author_profile(
            surname,
            sample.get("initials") or "",
            sample.get("given") or "",
            api_key,
        )

    profiles: dict[str, dict] = {}
    with ThreadPoolExecutor(max_workers=5) as pool:
        futures = {pool.submit(lookup, key): key for key in groups}
        for fut in as_completed(futures):
            key = futures[fut]
            try:
                profile = fut.result() or {}
            except Exception:
                profile = {}
            if profile.get("authid"):
                profiles[key] = profile
                for author in groups[key]:
                    apply_author_profile(author, profile)

    authids = sorted({(p.get("authid") or "") for p in profiles.values() if p.get("authid")})
    if not authids:
        return 0
    metrics: dict[str, dict] = {}
    with ThreadPoolExecutor(max_workers=5) as pool:
        futures = {pool.submit(fetch_author_metrics, authid, api_key): authid for authid in authids}
        for fut in as_completed(futures):
            authid = futures[fut]
            try:
                data = fut.result() or {}
            except Exception:
                data = {}
            if data:
                metrics[authid] = data
    if metrics:
        for rec in records:
            for author in rec.get("authors") or []:
                extra = metrics.get((author.get("authid") or "").strip())
                if extra:
                    apply_author_profile(author, extra)
    return len({p.get("authid") for p in profiles.values() if p.get("authid")})


def fetch_scopus_data(query: str, api_key: str, max_results: int | None) -> list[dict]:
    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    records = []
    start = 0
    page_size = 25
    total = None
    mode = "complete"
    while True:
        params = {"query": query, "count": page_size, "start": start}
        if mode == "complete":
            params["view"] = "COMPLETE"
        elif mode == "fields":
            params["field"] = SEARCH_FIELDS
        response = _scopus_get(headers, params)
        if mode == "complete" and response.status_code in {400, 401, 403}:
            mode = "fields"
            params.pop("view", None)
            params["field"] = SEARCH_FIELDS
            response = _scopus_get(headers, params)
        if mode == "fields" and response.status_code in {400, 401, 403}:
            mode = "standard"
            params.pop("field", None)
            response = _scopus_get(headers, params)
        if response.status_code != 200:
            raise RuntimeError(response.text)
        payload = response.json()
        search_results = payload.get("search-results") or {}
        if total is None:
            total = int(search_results.get("opensearch:totalResults", 0))
        entries = search_results.get("entry") or []
        if entries and isinstance(entries, dict):
            entries = [entries]
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            if entry.get("error"):
                continue
            if query_targets_gasu(query) and not entry_belongs_to_gasu(entry):
                continue
            cover_date = (entry.get("prism:coverDate") or "").strip()
            records.append(
                {
                    "title": (entry.get("dc:title") or "").strip(),
                    "journal": (entry.get("prism:publicationName") or "").strip(),
                    "year": cover_date[:4],
                    "cover_date": cover_date,
                    "volume": (entry.get("prism:volume") or "").strip(),
                    "issue": (entry.get("prism:issueIdentifier") or "").strip(),
                    "pages": (entry.get("prism:pageRange") or "").strip(),
                    "doi": (entry.get("prism:doi") or "").strip(),
                    "scopus_id": (entry.get("dc:identifier") or "").replace("SCOPUS_ID:", ""),
                    "eid": (entry.get("eid") or "").strip(),
                    "authors": parse_authors(entry),
                    "affiliation": format_affiliations(entry),
                    "issns": extract_issns(entry),
                }
            )
            if max_results and len(records) >= max_results:
                break
        start += page_size
        if start >= total or not entries:
            break
        if max_results and len(records) >= max_results:
            break
    records.sort(key=lambda item: item.get("cover_date") or "", reverse=True)
    return records


def expand_rsf_candidates(candidates: list[dict], window, api_key: str) -> tuple[list[dict], int, int]:
    expanded = []
    failed = 0
    counted = 0
    totals: dict[str, int] = {}
    for cand in candidates:
        row = dict(cand)
        authid = (cand.get("authid") or "").strip()
        if not authid:
            expanded.append(row)
            continue
        try:
            if authid not in totals:
                query = author_id_query(authid, window.from_year, window.to_year)
                totals[authid] = fetch_scopus_total(query, api_key)
            apply_author_total(row, totals[authid], "профиль Scopus (AU-ID)")
            counted += 1
        except Exception:
            failed += 1
        expanded.append(row)
    return expanded, failed, counted


def records_to_dataframe(records: list[dict]) -> pd.DataFrame:
    rows = []
    for rec in records:
        rows.append(
            {
                "Год": rec["year"],
                "Название": rec["title"],
                "Журнал": rec["journal"],
                "Авторы": format_authors_gost(rec["authors"]),
                "Author ID": "; ".join(author_ids(rec.get("authors"))),
                "Организации": rec.get("affiliation", ""),
                "Область знаний": rec.get("subject_areas") or "",
                "Квартиль SCImago": format_quartile_cell(rec),
                "SJR": rec.get("scimago_sjr") if rec.get("scimago_sjr") != "" else "",
                "Год SJR": rec.get("scimago_year") or "",
                "DOI": rec["doi"],
                "Scopus ID статьи": rec["scopus_id"],
            }
        )
    df = pd.DataFrame(rows)
    df.index = range(1, len(df) + 1)
    return df


def sort_records_for_bibliography(records: list[dict], date_filter: dict | None) -> list[dict]:
    def author_key(rec: dict) -> str:
        authors = rec.get("authors") or []
        if authors:
            surname = (authors[0].get("surname") or "").strip().lower()
            if surname:
                return surname
        return format_authors_gost(authors).lower()

    def year_key(rec: dict) -> int:
        year = rec.get("year") or ""
        return int(year) if year.isdigit() else 0

    if date_filter and date_filter.get("mode") == "range":
        return sorted(records, key=lambda rec: (author_key(rec), year_key(rec)))
    return sorted(records, key=author_key)


def build_docx(records: list[dict], fmt: str) -> BytesIO:
    doc = Document()
    title = "Список публикаций"
    doc.add_heading(title, level=1)
    for idx, rec in enumerate(records, start=1):
        text = format_gost(rec) if fmt == "ГОСТ 7.0.5" else format_apa(rec)
        doc.add_paragraph(f"{idx}. {text}")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_xlsx(
    records: list[dict],
    report_records: list[dict] | None = None,
    *,
    university: bool = False,
    grant_rows: list[dict] | None = None,
) -> BytesIO:
    df = records_to_dataframe(records)
    df["ГОСТ 7.0.5"] = [format_gost(r) for r in records]
    df["APA 7th"] = [format_apa(r) for r in records]
    source = report_records if report_records is not None else records
    report = build_report(source)
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Scopus")
        if report.year_rows:
            pd.DataFrame(report.year_rows).to_excel(writer, index=False, sheet_name="Динамика")
        if university:
            authors = top_authors(source, 20)
            if authors:
                pd.DataFrame(authors).to_excel(writer, index=False, sheet_name="Авторы")
        if grant_rows:
            pd.DataFrame(grant_rows).to_excel(writer, index=False, sheet_name="РНФ")
        if report.top_journals:
            pd.DataFrame(report.top_journals, columns=["Источник", "Публикаций"]).to_excel(
                writer, index=False, sheet_name="Источники"
            )
        area_rows = area_share_rows(source)
        if area_rows:
            pd.DataFrame(area_rows).to_excel(writer, index=False, sheet_name="Области знаний")
        q_rows = quartile_share_rows(source)
        if q_rows:
            pd.DataFrame(q_rows).to_excel(writer, index=False, sheet_name="Квартили")
    buf.seek(0)
    return buf


def render_report_block(
    report: ReportData,
    records: list[dict],
    *,
    author_mode: bool,
    precise_author: bool,
    university: bool,
    author_last: str = "",
    show_top_authors: bool = True,
) -> None:
    st.subheader("Динамика для отчёта")
    if author_mode and not precise_author:
        st.warning(
            "Профиль Scopus не сопоставился однозначно: показаны работы по фамилии. "
            "Возможны однофамильцы. Для персонального отчёта укажите ORCID."
        )
    st.write(report_sentence(report))
    if report.counts:
        png = report_chart_png(report)
        st.image(png, width=560)
        st.download_button(
            "Скачать график (PNG)",
            data=png,
            file_name="scopus_dinamika.png",
            mime="image/png",
            key="download_year_png",
        )
        st.caption("По году публикации. Годы без работ показаны как 0. Год к году — в таблице ниже.")
    if report.year_rows:
        st.dataframe(pd.DataFrame(report.year_rows), hide_index=True, use_container_width=True)

    if records:
        scope = report_scope_label(
            records,
            university=university,
            author_last=author_last,
        )
        area_rows = area_share_rows(records)
        q_rows = quartile_share_rows(records)
        st.markdown("**Состав выборки**")
        st.caption(
            f"Круги — области знаний и квартили всех статей текущего поиска "
            f"({scope}, {ru_publications(report.total)}). "
            "Период задаётся при поиске: текущий год, последние 5 лет, свой диапазон или с 1993 года."
        )
        png = report_area_png(
            grouped_area_counts(area_rows),
            title=f"{scope} · области знаний",
        )
        png_q = report_quartile_png(q_rows, title=f"{scope} · квартили SCImago")
        left, right = st.columns(2, gap="medium")
        with left:
            st.image(png, use_container_width=True)
        with right:
            st.image(png_q, use_container_width=True)
        dl_left, dl_right = st.columns(2, gap="medium")
        with dl_left:
            st.download_button(
                "Скачать области (PNG)",
                data=png,
                file_name="scopus_oblasti.png",
                mime="image/png",
                key="download_area_png",
                use_container_width=True,
            )
        with dl_right:
            st.download_button(
                "Скачать квартили (PNG)",
                data=png_q,
                file_name="scopus_kvartili.png",
                mime="image/png",
                key="download_quartile_png",
                use_container_width=True,
            )
        st.caption(
            "Слева — основная область журнала в Scopus (одна статья — один сектор; "
            "без ISSN — «Не указано»). Справа — лучший квартиль журнала SCImago; "
            "если года статьи ещё нет в рейтинге, берётся последний закрытый год."
        )
        if report.total:
            detail_left, detail_right = st.columns(2, gap="medium")
            with detail_left:
                if area_rows:
                    st.dataframe(pd.DataFrame(area_rows), hide_index=True, use_container_width=True)
            with detail_right:
                if q_rows:
                    st.dataframe(pd.DataFrame(q_rows), hide_index=True, use_container_width=True)

    if university and show_top_authors:
        st.markdown("**Наиболее активные авторы**")
        top_n = st.radio(
            "Размер списка",
            [5, 10, 20],
            index=1,
            horizontal=True,
            format_func=lambda n: f"Топ {n}",
            key="top_authors_n",
            label_visibility="collapsed",
        )
        author_rows = top_authors(records, int(top_n))
        if author_rows:
            st.dataframe(
                pd.DataFrame(author_rows),
                hide_index=True,
                use_container_width=True,
            )
            st.caption(
                f"{report_scope_label(records, university=True)}. "
                "Соавторство считается: статья входит в показатель каждого автора. "
                "Q1–Q4 — квартиль журнала этой статьи. "
                "Author ID, ORCID, h-индекс, документы и цитирования — из профиля Scopus, не из среза."
            )

    if report.top_journals:
        st.caption("Топ источников")
        st.dataframe(
            pd.DataFrame(report.top_journals, columns=["Источник", "Публикаций"]),
            hide_index=True,
            use_container_width=True,
        )
    if report.external_share is not None:
        st.caption(
            f"Работ с внешней организацией в аффилиации: {report.external_count} из {report.total} "
            f"({report.external_share}%)."
        )


st.set_page_config(page_title="Мониторинг публикаций Scopus", layout="wide")
auth_cookies = cookie_manager()
if not require_login(auth_cookies):
    st.stop()

if st.session_state.get("records_version") != APP_VERSION:
    for key in (
        "records",
        "query",
        "date_filter",
        "dynamics_records",
        "show_report",
        "search_mode",
        "author_last",
        "author_orcid",
        "only_gasu",
        "issn_subject_cache",
        "area_chart_year",
        "top_authors_n",
        "grant_min",
        "grant_contest_year",
        "grant_from_year",
        "grant_rows",
        "grant_failed",
        "grant_with_id",
        "grant_people",
        "grant_named",
    ):
        st.session_state.pop(key, None)

st.title("Мониторинг публикаций Scopus")

api_key = load_api_key()
on_cloud = running_on_streamlit_cloud()
with st.sidebar:
    st.header("Доступ")
    if st.button("Выйти"):
        logout(auth_cookies)
    st.header("Scopus")
    if not api_key:
        if on_cloud:
            st.warning(
                "Добавьте `SCOPUS_API_KEY` в Settings → Secrets приложения. "
                "На Cloud файл `.env` не сохраняется после перезапуска."
            )
        else:
            st.write("Введите API-ключ один раз. Он сохранится локально в `.env`.")
            key_input = st.text_input("API-ключ Scopus", type="password")
            if st.button("Сохранить ключ"):
                if key_input.strip():
                    save_api_key(key_input)
                    st.success("Ключ сохранен. Перезагружаю...")
                    st.rerun()
                else:
                    st.warning("Введите корректный ключ.")
    else:
        st.success("API-ключ найден.")
        st.caption("Ключ берётся из Secrets (Cloud) или из `.env` (локально). Пользователи его не вводят.")
        built = lookup_built_on()
        nxt = next_refresh_date()
        max_year = load_meta().get("max_year")
        st.caption(
            f"Справочник квартилей SCImago обновлён {format_ru_date(built)}. "
            f"Следующее автообновление: {format_ru_date(nxt)}."
        )
        if max_year:
            st.caption(f"В файле рейтинги журналов по {max_year} год включительно.")

st.caption(
    "Мониторинг ГАГУ включает статьи, где университет указан хотя бы как одна из "
    "аффилиаций (в том числе вместе с другими организациями)."
)

window = rsf_window()
quick1, quick2, quick3 = st.columns(3)
with quick1:
    quick_check = st.button(
        "Статьи ГАГУ за текущий год",
        type="primary",
        use_container_width=True,
        key="quick_year",
    )
with quick2:
    quick_rsf8 = st.button(
        f"РНФ {window.contest_year}, от 8 статей",
        use_container_width=True,
        key="quick_rsf8",
    )
with quick3:
    quick_rsf5 = st.button(
        f"РНФ {window.contest_year}, от 5 статей",
        use_container_width=True,
        key="quick_rsf5",
    )
st.caption(
    f"РНФ: ближайший конкурс — {window.contest_year} год. "
    f"Порог считается по всем статьям Scopus автора с {window.from_label}, а не только с аффилиацией ГАГУ. "
    "В список попадают люди, у которых в этом окне есть хотя бы одна статья с ГАГУ. "
    "Штат или совместительство Scopus не показывает."
)

mode = st.radio("Режим поиска", ["Мониторинг ГАГУ", "Поиск по автору"], horizontal=True)

author_last = ""
author_orcid = ""
only_gasu = False
if mode == "Поиск по автору":
    author_last = st.text_input("Фамилия автора")
    author_orcid = st.text_input("ORCID (если есть)")
    only_gasu = st.checkbox("Только аффилиация ГАГУ", value=False)

time_filter = st.radio(
    "Период",
    ["Текущий год", "Последние 5 лет", "Диапазон лет", "С 1993 года"],
    horizontal=True,
)
start_year = None
end_year = None
if time_filter == "Диапазон лет":
    col1, col2 = st.columns(2)
    with col1:
        start_year = st.number_input("С", min_value=1900, max_value=2100, value=2020, step=1)
    with col2:
        end_year = st.number_input("По", min_value=1900, max_value=2100, value=datetime.now().year, step=1)

search_clicked = st.button("Найти публикации")
grant_min = None

if quick_check:
    mode = "Мониторинг ГАГУ"
    time_filter = "Текущий год"
    search_clicked = True
elif quick_rsf8 or quick_rsf5:
    mode = "Мониторинг ГАГУ"
    time_filter = "РНФ"
    grant_min = 8 if quick_rsf8 else 5
    search_clicked = True

date_filter = None
if time_filter == "Текущий год":
    date_filter = make_date_filter("current", None, None)
elif time_filter == "Последние 5 лет":
    date_filter = make_date_filter("last5", None, None)
elif time_filter == "С 1993 года":
    date_filter = make_date_filter("since_founding", None, None)
elif time_filter == "РНФ":
    date_filter = {
        "mode": "range",
        "year": window.from_year,
        "year_start": window.from_year,
        "year_end": window.to_year,
    }
else:
    date_filter = make_date_filter("range", int(start_year), int(end_year))

if search_clicked:
    if not api_key:
        st.error("Нужен API-ключ Scopus. Введите его в боковой панели.")
        st.stop()
    if mode == "Поиск по автору" and not author_orcid and not author_last:
        st.error("Для поиска по автору укажите фамилию или ORCID.")
        st.stop()
    identity = "university"
    if mode == "Поиск по автору":
        if author_orcid:
            identity = "orcid"
            query = build_query(mode, author_last, author_orcid, date_filter, only_gasu)
        else:
            with st.spinner("Ищем профиль автора в Scopus..."):
                authid = resolve_scopus_authid(author_last, "", "", api_key)
            if authid:
                identity = "au-id"
                query = author_papers_query(authid, date_filter, only_gasu)
            else:
                identity = "surname"
                query = build_query(mode, author_last, author_orcid, date_filter, only_gasu)
    else:
        query = build_query(mode, author_last, author_orcid, date_filter, only_gasu)
    with st.spinner("Идет поиск в Scopus..."):
        try:
            records = fetch_scopus_data(query, api_key, None)
        except Exception as exc:
            st.error("Ошибка запроса к Scopus API.")
            st.code(str(exc))
            st.stop()
    if records and grant_min:
        records = [rec for rec in records if record_in_rsf_window(rec, window)]
        for rec in records:
            rec.setdefault("subject_abbrevs", [])
            rec.setdefault("subject_areas", "Не указано")
        try:
            attach_scimago(records)
        except Exception:
            for rec in records:
                rec.setdefault("scimago_quartile", "Нет")
    elif records:
        with st.spinner("Определяем области знаний по журналам Scopus..."):
            try:
                fill_subject_areas(records, api_key)
            except Exception:
                for rec in records:
                    rec.setdefault("subject_abbrevs", [])
                    rec.setdefault("subject_areas", "Не указано")
        with st.spinner("Сопоставляем квартили SCImago по ISSN журнала..."):
            try:
                attach_scimago(records)
            except Exception:
                for rec in records:
                    rec.setdefault("scimago_quartile", "Нет")

    if not records:
        st.info("Статей по данному запросу не найдено.")
        st.stop()

    with st.spinner("Дополняем авторов по карточкам статей, если API это позволяет..."):
        enrich_record_authors(records, api_key)
    with st.spinner("Загружаем профили Scopus: Author ID, ORCID, h-индекс..."):
        stamp_author_profiles(records, api_key)

    st.session_state["records"] = records
    st.session_state["date_filter"] = date_filter
    st.session_state["query"] = query
    st.session_state["records_version"] = APP_VERSION
    st.session_state["search_mode"] = mode
    st.session_state["author_last"] = author_last
    st.session_state["author_orcid"] = author_orcid
    st.session_state["only_gasu"] = only_gasu
    st.session_state["author_identity"] = identity
    st.session_state["show_report"] = False
    st.session_state.pop("dynamics_records", None)
    if grant_min:
        st.session_state["grant_min"] = grant_min
        st.session_state["grant_contest_year"] = window.contest_year
        st.session_state["grant_from_year"] = window.from_year
        candidates = rsf_candidates(records)
        st.session_state["grant_rows"] = rsf_eligibility_rows(candidates, grant_min)
        st.session_state["grant_people"] = len(candidates)
        st.session_state["grant_with_id"] = sum(1 for cand in candidates if cand.get("authid"))
        st.session_state["grant_failed"] = 0
        st.session_state["grant_named"] = 0
        with st.spinner("Считаем все статьи Scopus по Author ID..."):
            expanded, failed, counted = expand_rsf_candidates(candidates, window, api_key)
        st.session_state["grant_rows"] = rsf_eligibility_rows(expanded, grant_min)
        st.session_state["grant_failed"] = failed
        st.session_state["grant_named"] = counted
        st.session_state["grant_with_id"] = sum(1 for cand in expanded if cand.get("authid"))
    else:
        st.session_state.pop("grant_min", None)
        st.session_state.pop("grant_contest_year", None)
        st.session_state.pop("grant_from_year", None)
        st.session_state.pop("grant_rows", None)
        st.session_state.pop("grant_failed", None)
        st.session_state.pop("grant_with_id", None)
        st.session_state.pop("grant_people", None)
        st.session_state.pop("grant_named", None)

if "records" in st.session_state and st.session_state["records"]:
    records = st.session_state["records"]
    active_date_filter = st.session_state.get("date_filter")
    records_for_list = sort_records_for_bibliography(records, active_date_filter)
    grant_min_saved = st.session_state.get("grant_min")
    grant_rows: list[dict] = list(st.session_state.get("grant_rows") or [])
    if grant_min_saved:
        contest = st.session_state.get("grant_contest_year")
        from_year = st.session_state.get("grant_from_year")
        st.subheader("Потенциальные грантодержатели РНФ")
        st.write(
            f"Конкурс {contest}: порог — не менее {grant_min_saved} публикаций Scopus "
            f"с января {from_year}, с любой аффилиацией."
        )
        g1, g2, g3 = st.columns(3)
        g1.metric("Могут подавать сейчас", len(grant_rows))
        g2.metric("Статей ГАГУ в окне", len(records))
        g3.metric("Порог", f"{grant_min_saved} Scopus")
        if grant_rows:
            st.dataframe(pd.DataFrame(grant_rows), hide_index=True, use_container_width=True)
        else:
            st.info("Никто не набрал порог по всем статьям Scopus в этом окне.")
        failed = int(st.session_state.get("grant_failed") or 0)
        with_id = int(st.session_state.get("grant_with_id") or 0)
        people = int(st.session_state.get("grant_people") or 0)
        named = int(st.session_state.get("grant_named") or 0)
        counted = with_id
        if people and counted == 0:
            st.warning(
                "Профили Scopus не сопоставились (Author Search не вернул однозначный Author ID). "
                "Порог посчитан по тем же статьям ГАГУ, что в списке ниже."
            )
        st.caption(
            "Счёт РНФ — по тем же статьям, что список литературы ниже. "
            "Профиль автора (Author ID, ORCID, h-индекс, документы, цитирования) берётся из Author Search / Author Retrieval, "
            "как на странице автора в Scopus. Если API не отдаёт поле — ячейка пустая. "
            "«Всего Scopus» в окне РНФ уточняется запросом AU-ID. "
            "Это оценка для мониторинга, не экспертиза заявки."
            + (
                f" Author ID есть у {counted} из {people} авторов."
                if people
                else ""
            )
            + (f" Ошибки запроса: {failed}." if failed else "")
        )

    st.subheader("Результаты")
    slice_report = build_report(records)
    kpi1, kpi2, kpi3 = st.columns(3)
    kpi1.metric("Публикаций в срезе", slice_report.total)
    kpi2.metric("Охват лет", slice_report.year_label)
    kpi3.metric("Журналов и изданий", slice_report.unique_journals)
    st.caption(
        "В список входят только записи, где в ответе Scopus видно ГАГУ по названию или городу "
        "(в том числе как одна из нескольких аффилиаций)."
    )

    active_filter = st.session_state.get("date_filter") or {}
    filter_span = 1
    if active_filter.get("year_start") and active_filter.get("year_end"):
        filter_span = int(active_filter["year_end"]) - int(active_filter["year_start"]) + 1
    can_chart_from_slice = slice_report.distinct_years >= 2 or filter_span >= 2

    if not st.session_state.get("show_report"):
        if can_chart_from_slice:
            if st.button("Показать динамику для отчёта"):
                st.session_state["show_report"] = True
                st.rerun()
        elif st.button("Динамика за 5 лет"):
            if not api_key:
                st.error("Нужен API-ключ Scopus.")
            else:
                last5 = make_date_filter("last5", None, None)
                dyn_query = build_query(
                    st.session_state.get("search_mode") or "Мониторинг ГАГУ",
                    st.session_state.get("author_last") or "",
                    st.session_state.get("author_orcid") or "",
                    last5,
                    bool(st.session_state.get("only_gasu")),
                )
                with st.spinner("Строим динамику за 5 лет..."):
                    try:
                        dyn_records = fetch_scopus_data(dyn_query, api_key, None)
                        try:
                            fill_subject_areas(dyn_records, api_key)
                        except Exception:
                            for rec in dyn_records:
                                rec.setdefault("subject_abbrevs", [])
                                rec.setdefault("subject_areas", "Не указано")
                        try:
                            attach_scimago(dyn_records)
                        except Exception:
                            for rec in dyn_records:
                                rec.setdefault("scimago_quartile", "Нет")
                        st.session_state["dynamics_records"] = dyn_records
                        st.session_state["show_report"] = True
                        st.rerun()
                    except Exception as exc:
                        st.error("Не удалось получить данные за 5 лет.")
                        st.code(str(exc))

    if st.session_state.get("show_report"):
        report_source = st.session_state.get("dynamics_records")
        if report_source is None and can_chart_from_slice:
            report_source = records
        if report_source:
            render_report_block(
                build_report(report_source),
                report_source,
                author_mode=(st.session_state.get("search_mode") == "Поиск по автору"),
                precise_author=st.session_state.get("author_identity") in {"orcid", "au-id"},
                university=(st.session_state.get("search_mode") != "Поиск по автору"),
                author_last=st.session_state.get("author_last") or "",
                show_top_authors=True,
            )
        else:
            st.info("За этот период публикаций не найдено.")

    df = records_to_dataframe(records)
    st.dataframe(df, use_container_width=True)
    st.caption(
        "Scopus ID статьи — документ. Author ID — профиль автора в Scopus; по нему считается полный список работ. "
        "Квартиль — лучший квартиль журнала в SCImago. Если у статьи год новее рейтинга, в скобках указан год SJR "
        f"(сейчас до {load_meta().get('max_year') or '—'}). Это не оценка текста статьи."
    )
    if st.session_state.get("query"):
        with st.expander("Запрос Scopus"):
            st.code(st.session_state["query"])

    st.subheader("Готовый список литературы")
    format_choice = st.selectbox("Формат", ["ГОСТ 7.0.5", "APA 7th Edition"])
    formatted_list = [
        format_gost(rec) if format_choice == "ГОСТ 7.0.5" else format_apa(rec)
        for rec in records_for_list
    ]
    st.markdown("\n".join([f"{i}. {text}" for i, text in enumerate(formatted_list, start=1)]))

    docx_buffer = build_docx(records_for_list, format_choice)
    xlsx_buffer = build_xlsx(
        records_for_list,
        st.session_state.get("dynamics_records") or records,
        university=(st.session_state.get("search_mode") != "Поиск по автору"),
        grant_rows=grant_rows or None,
    )

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "Скачать .docx",
            data=docx_buffer,
            file_name="scopus_publications.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    with col2:
        st.download_button(
            "Скачать .xlsx",
            data=xlsx_buffer,
            file_name="scopus_publications.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

st.markdown("---")
st.caption("© Алексеев П.В., pavel.alekseev.gasu@gmail.com, Горно-Алтайский государственный университет")
st.caption(f"Версия {APP_VERSION} · последнее обновление {last_updated_label()}")
